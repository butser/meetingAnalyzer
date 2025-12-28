"""
SRS document generator module
"""

import os
from datetime import datetime
from typing import Dict, List
from pathlib import Path
import json


class SRSGenerator:
    """Generates Software Requirements Specification documents"""
    
    def __init__(self, output_dir: str = "output"):
        """
        Initialize the SRS generator
        
        Args:
            output_dir: Directory to save generated documents
        """
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)
    
    def generate_markdown(self, requirements: Dict, project_name: str = "Project") -> str:
        """
        Generate SRS document in Markdown format
        
        Args:
            requirements: Requirements data
            project_name: Name of the project
            
        Returns:
            Path to generated Markdown file
        """
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
        # Start building the markdown content
        content = f"""# Software Requirements Specification (SRS)

**Project:** {project_name}  
**Date Generated:** {timestamp}  
**Version:** 1.0

---

## Table of Contents

1. [Introduction](#introduction)
2. [Project Overview](#project-overview)
3. [Functional Requirements](#functional-requirements)
4. [Non-Functional Requirements](#non-functional-requirements)
5. [Technical Requirements](#technical-requirements)
6. [UI/UX Requirements](#ui-ux-requirements)
7. [Issues and Concerns](#issues-and-concerns)

---

## 1. Introduction

This Software Requirements Specification (SRS) document was automatically generated from meeting recordings using AI-powered analysis. It combines audio transcription and visual content analysis to extract comprehensive project requirements.

---

"""
        
        # Add sections based on requirements data
        if isinstance(requirements, dict):
            # Project Overview
            content += "## 2. Project Overview\n\n"
            if "PROJECT OVERVIEW" in requirements or "project_overview" in requirements:
                overview = requirements.get("PROJECT OVERVIEW", requirements.get("project_overview", {}))
                if isinstance(overview, dict):
                    for key, value in overview.items():
                        content += f"**{key.replace('_', ' ').title()}:** {value}\n\n"
                else:
                    content += f"{overview}\n\n"
            else:
                content += "*To be determined based on meeting analysis.*\n\n"
            
            # Functional Requirements
            content += "## 3. Functional Requirements\n\n"
            if "FUNCTIONAL REQUIREMENTS" in requirements or "functional_requirements" in requirements:
                func_reqs = requirements.get("FUNCTIONAL REQUIREMENTS", requirements.get("functional_requirements", []))
                if isinstance(func_reqs, list):
                    for i, req in enumerate(func_reqs, 1):
                        content += f"### FR-{i:03d}\n\n{req}\n\n"
                elif isinstance(func_reqs, dict):
                    for key, value in func_reqs.items():
                        content += f"### {key}\n\n{value}\n\n"
                else:
                    content += f"{func_reqs}\n\n"
            else:
                content += "*Functional requirements will be extracted from meeting analysis.*\n\n"
            
            # Non-Functional Requirements
            content += "## 4. Non-Functional Requirements\n\n"
            if "NON-FUNCTIONAL REQUIREMENTS" in requirements or "non_functional_requirements" in requirements:
                nfr = requirements.get("NON-FUNCTIONAL REQUIREMENTS", requirements.get("non_functional_requirements", {}))
                if isinstance(nfr, dict):
                    for key, value in nfr.items():
                        content += f"### {key.replace('_', ' ').title()}\n\n{value}\n\n"
                else:
                    content += f"{nfr}\n\n"
            else:
                content += "*Non-functional requirements will be extracted from meeting analysis.*\n\n"
            
            # Technical Requirements
            content += "## 5. Technical Requirements\n\n"
            if "TECHNICAL REQUIREMENTS" in requirements or "technical_requirements" in requirements:
                tech_reqs = requirements.get("TECHNICAL REQUIREMENTS", requirements.get("technical_requirements", {}))
                if isinstance(tech_reqs, dict):
                    for key, value in tech_reqs.items():
                        content += f"**{key.replace('_', ' ').title()}:** {value}\n\n"
                else:
                    content += f"{tech_reqs}\n\n"
            else:
                content += "*Technical requirements will be extracted from meeting analysis.*\n\n"
            
            # UI/UX Requirements
            content += "## 6. UI/UX Requirements\n\n"
            if "UI/UX REQUIREMENTS" in requirements or "ui_ux_requirements" in requirements:
                ui_reqs = requirements.get("UI/UX REQUIREMENTS", requirements.get("ui_ux_requirements", ""))
                content += f"{ui_reqs}\n\n"
            else:
                content += "*UI/UX requirements will be extracted from visual analysis.*\n\n"
            
            # Issues and Concerns
            content += "## 7. Issues and Concerns\n\n"
            if "ISSUES AND CONCERNS" in requirements or "issues_and_concerns" in requirements:
                issues = requirements.get("ISSUES AND CONCERNS", requirements.get("issues_and_concerns", []))
                if isinstance(issues, list):
                    for issue in issues:
                        content += f"- {issue}\n"
                else:
                    content += f"{issues}\n"
                content += "\n"
            else:
                content += "*Issues and concerns will be extracted from meeting discussion.*\n\n"
            
            # Add raw analysis if present
            if "raw_analysis" in requirements:
                content += "---\n\n## Raw Analysis\n\n"
                content += f"{requirements['raw_analysis']}\n\n"
        
        else:
            content += f"\n{requirements}\n\n"
        
        # Add appendix
        content += """---

## Appendix

### Document Information

- **Generated By:** Meeting Analyzer Tool
- **Source:** Automated analysis of meeting video and audio
- **Analysis Method:** AI-powered transcription and visual content analysis

### Notes

This document was automatically generated and should be reviewed and refined by the project team.
"""
        
        # Save to file
        output_path = os.path.join(self.output_dir, f"SRS_{project_name.replace(' ', '_')}.md")
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(content)
        
        print(f"SRS document (Markdown) generated: {output_path}")
        return output_path
    
    def generate_docx(self, requirements: Dict, project_name: str = "Project") -> str:
        """
        Generate SRS document in DOCX format
        
        Args:
            requirements: Requirements data
            project_name: Name of the project
            
        Returns:
            Path to generated DOCX file
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            # Title
            title = doc.add_heading('Software Requirements Specification (SRS)', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Project info
            doc.add_paragraph(f'Project: {project_name}')
            doc.add_paragraph(f'Date: {datetime.now().strftime("%Y-%m-%d")}')
            doc.add_paragraph(f'Version: 1.0')
            
            doc.add_page_break()
            
            # Add content sections
            sections = [
                ("Introduction", "introduction"),
                ("Project Overview", "project_overview"),
                ("Functional Requirements", "functional_requirements"),
                ("Non-Functional Requirements", "non_functional_requirements"),
                ("Technical Requirements", "technical_requirements"),
                ("UI/UX Requirements", "ui_ux_requirements"),
                ("Issues and Concerns", "issues_and_concerns")
            ]
            
            for section_title, section_key in sections:
                doc.add_heading(section_title, 1)
                
                if isinstance(requirements, dict) and section_key in requirements:
                    data = requirements[section_key]
                    if isinstance(data, dict):
                        for key, value in data.items():
                            doc.add_heading(key.replace('_', ' ').title(), 2)
                            doc.add_paragraph(str(value))
                    elif isinstance(data, list):
                        for item in data:
                            doc.add_paragraph(str(item), style='List Bullet')
                    else:
                        doc.add_paragraph(str(data))
                else:
                    doc.add_paragraph(f"{section_title} will be extracted from meeting analysis.")
                
                doc.add_paragraph()
            
            # Save document
            output_path = os.path.join(self.output_dir, f"SRS_{project_name.replace(' ', '_')}.docx")
            doc.save(output_path)
            
            print(f"SRS document (DOCX) generated: {output_path}")
            return output_path
            
        except ImportError:
            print("python-docx not installed, skipping DOCX generation")
            return None
    
    def save_json(self, requirements: Dict, project_name: str = "Project") -> str:
        """
        Save requirements as JSON for further processing
        
        Args:
            requirements: Requirements data
            project_name: Name of the project
            
        Returns:
            Path to JSON file
        """
        output_path = os.path.join(self.output_dir, f"requirements_{project_name.replace(' ', '_')}.json")
        
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(requirements, f, indent=2, ensure_ascii=False)
        
        print(f"Requirements JSON saved: {output_path}")
        return output_path
